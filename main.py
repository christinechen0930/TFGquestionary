import streamlit as st
import requests
import os
import torch
from sentence_transformers import SentenceTransformer, util
import fitz  # PyMuPDF
import docx  # 需要安裝 python-docx
from tavily import TavilyClient

# ====== 設定 API Key ======
TAVILY_API_KEY = "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"
GEMINI_API_KEY = "AIzaSyC25eTdPDzuMqv3ZE_I8l6gpuv0faBA88c"

tavily_client = TavilyClient(api_key=TAVILY_API_KEY)

# ====== 建立必要資料夾 ======
DOWNLOAD_FOLDER = "downloads"
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

# ====== 載入語義模型 ======
MODEL_PATH = "./paraphrase-MiniLM-L6-v2"
if not os.path.exists(MODEL_PATH):
    st.error("❌ 找不到語義搜尋模型，請先使用以下指令下載：\n"
             "`git clone https://huggingface.co/sentence-transformers/paraphrase-MiniLM-L6-v2`")
    st.stop()

model = SentenceTransformer(MODEL_PATH)

# ====== 功能區塊 ======

def search_and_download_pdfs(keyword):
    query = f"site:fg.tp.edu.tw {keyword} filetype:pdf"
    response = tavily_client.search(query)
    pdf_links = [result["url"] for result in response.get("results", []) if result["url"].endswith(".pdf")]

    if not pdf_links:
        return "❌ 沒有找到相關的 PDF 檔案！"

    pdf_paths = []
    for index, pdf_url in enumerate(pdf_links):
        try:
            response = requests.get(pdf_url, timeout=10)
            pdf_filename = os.path.join(DOWNLOAD_FOLDER, f"{keyword}_{index + 1}.pdf")
            with open(pdf_filename, "wb") as f:
                f.write(response.content)
            pdf_paths.append(pdf_filename)
        except Exception as e:
            return f"❌ 下載失敗：{pdf_url}，錯誤：{e}"

    return pdf_paths

def read_pdf(file_path):
    try:
        doc = fitz.Document(file_path)
        return [page.get_text() for page in doc]
    except Exception as e:
        return [f"讀取 PDF 錯誤：{str(e)}"]

def read_docx(file):
    try:
        doc = docx.Document(file)
        return [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    except Exception as e:
        return [f"讀取 DOCX 錯誤：{str(e)}"]

def retrieve_relevant_content(task, paragraphs):
    paragraph_embeddings = model.encode(paragraphs, convert_to_tensor=True)
    query_embedding = model.encode(task, convert_to_tensor=True)
    scores = util.pytorch_cos_sim(query_embedding, paragraph_embeddings)[0]
    top_k = min(5, len(paragraphs))
    top_results = torch.topk(scores, k=top_k)
    return " ".join([paragraphs[idx] for idx in top_results.indices])

def generate_response_combined(task, keyword, file):
    if file:
        if file.name.endswith(".pdf"):
            paragraphs = read_pdf(file)
        elif file.name.endswith(".docx"):
            paragraphs = read_docx(file)
        else:
            return "❌ 不支援的檔案格式，請上傳 PDF 或 DOCX"
    else:
        if not keyword.strip():
            return "❌ 請輸入關鍵字或上傳檔案"

        pdf_paths = search_and_download_pdfs(keyword)
        if isinstance(pdf_paths, str):
            return pdf_paths

        paragraphs = []
        for pdf_path in pdf_paths:
            paragraphs.extend(read_pdf(pdf_path))

    if not paragraphs or "錯誤" in paragraphs[0]:
        return paragraphs[0]

    relevant_content = retrieve_relevant_content(task, paragraphs)
    if not relevant_content.strip():
        return "❌ 找不到與問題相關的內容，請嘗試其他關鍵字。"

    prompt = f"""
    請根據以下文件內容回答問題：
    問題：{task}
    相關內容：
    {relevant_content}

    請用繁體中文回答，並用條列式或摘要方式簡潔表達。
    """

    api_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"parts": [{"text": prompt}]}]}

    response = requests.post(f"{api_url}?key={GEMINI_API_KEY}", json=payload, headers=headers)

    if response.status_code == 200:
        response_json = response.json()
        if "candidates" in response_json and len(response_json["candidates"]) > 0:
            return response_json["candidates"][0]["content"]["parts"][0]["text"]
        else:
            return "❌ 無法取得模型回答"
    else:
        return f"❌ 錯誤：{response.status_code}, {response.text}"

# ====== Streamlit UI ======

st.title("🌱 綠園事務詢問欄")

task = st.text_input("輸入詢問事項", "例如：如何申請交換學生？")
keyword = st.text_input("輸入關鍵字（自動搜尋北一女 PDF）", "例如：招生簡章")
file_input = st.file_uploader("或上傳 PDF / DOCX", type=["pdf", "docx"])

if st.button("生成回答"):
    response = generate_response_combined(task, keyword, file_input)
    st.markdown(response)
